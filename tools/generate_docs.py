#!/usr/bin/env python3
"""
Genera el documento Word completo de TravelBox Peru:
  - Diagramas de Casos de Uso (imágenes generadas con matplotlib)
  - Diagrama de Estados de Reserva
  - Diagrama de Flujo de Pagos
  - Diagrama de Arquitectura de Capas
  - Datos de Prueba Reales
"""
import os
import math
import textwrap
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
IMG_DIR = os.path.join(OUT_DIR, "diagrams")
os.makedirs(IMG_DIR, exist_ok=True)

# ── Styling helpers ──────────────────────────────────────────────────────────

BLUE   = "#1565C0"
DBLUE  = "#0D47A1"
LBLUE  = "#E3F2FD"
GREEN  = "#2E7D32"
LGREEN = "#E8F5E9"
ORANGE = "#E65100"
LORANGE= "#FFF3E0"
RED    = "#C62828"
LRED   = "#FFEBEE"
PURPLE = "#6A1B9A"
LPURPLE= "#F3E5F5"
GRAY   = "#616161"
LGRAY  = "#F5F5F5"
TEAL   = "#00695C"
LTEAL  = "#E0F2F1"
AMBER  = "#FF8F00"
LAMBER = "#FFF8E1"

def set_cell_bg(cell, color_hex):
    """Set background color for a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex.replace("#", ""))
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)

# ── DIAGRAM 1: Use Case - Turista ───────────────────────────────────────────

def draw_use_case_turista():
    fig, ax = plt.subplots(1, 1, figsize=(14, 11))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 11)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # Title
    ax.text(7, 10.5, "Diagrama de Casos de Uso — Actor: Turista",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    # Actor
    ax.plot(1.5, 5.5, 'o', markersize=20, color=BLUE, zorder=5)
    ax.plot([1.5, 1.5], [5.0, 4.2], color=BLUE, lw=2, zorder=5)
    ax.plot([1.0, 2.0], [4.8, 4.8], color=BLUE, lw=2, zorder=5)
    ax.plot([1.5, 1.1], [4.2, 3.5], color=BLUE, lw=2, zorder=5)
    ax.plot([1.5, 1.9], [4.2, 3.5], color=BLUE, lw=2, zorder=5)
    ax.text(1.5, 3.1, "Turista", fontsize=11, ha='center', fontweight='bold', color=DBLUE)

    # Use cases
    cases = [
        (7, 9.5,  "UC-01: Registrarse / Login", LBLUE, BLUE),
        (7, 8.5,  "UC-02: Login Social\n(Google / Microsoft)", LPURPLE, PURPLE),
        (7, 7.4,  "UC-03: Buscar Almacenes\nCercanos", LGREEN, GREEN),
        (7, 6.3,  "UC-04: Crear Reserva\nde Equipaje", LBLUE, BLUE),
        (7, 5.2,  "UC-05: Pagar con Tarjeta\n(Izipay Krypton)", LORANGE, ORANGE),
        (7, 4.1,  "UC-06: Pagar en Efectivo\n(Cash)", LAMBER, AMBER),
        (7, 3.0,  "UC-07: Ver Tracking\nde Reserva", LTEAL, TEAL),
        (7, 1.9,  "UC-08: Recibir\nNotificaciones", LRED, RED),
        (7, 0.8,  "UC-09: Calificar\nAlmacén", LPURPLE, PURPLE),
        (11.5, 9.5, "UC-10: Gestionar Perfil", LGRAY, GRAY),
        (11.5, 8.4, "UC-11: Cambiar Email", LGRAY, GRAY),
        (11.5, 7.3, "UC-12: Recuperar\nContraseña", LRED, RED),
        (11.5, 6.2, "UC-13: Reportar\nIncidente", LORANGE, ORANGE),
        (11.5, 5.1, "UC-14: Ver Historial\nde Pagos", LBLUE, BLUE),
        (11.5, 4.0, "UC-15: Cancelar\nReserva", LRED, RED),
    ]

    for (cx, cy, label, bg, border) in cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.8, 0.85, facecolor=bg,
                                    edgecolor=border, linewidth=1.8, zorder=3)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, fontsize=7.5, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.15)
        # Line from actor to case
        ax.annotate("", xy=(cx - 1.9, cy), xytext=(2.2, min(max(cy, 3.5), 5.5)),
                     arrowprops=dict(arrowstyle='->', color='#90A4AE', lw=1.0),
                     zorder=2)

    # System boundary
    rect = mpatches.FancyBboxPatch((4.5, 0.2), 9.2, 10.0, boxstyle="round,pad=0.3",
                                    facecolor='none', edgecolor=DBLUE, linewidth=2,
                                    linestyle='--', zorder=1)
    ax.add_patch(rect)
    ax.text(9.1, 0.45, "Sistema TravelBox Peru — Backend API", fontsize=9,
            ha='center', color=DBLUE, fontstyle='italic')

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "uc_turista.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 2: Use Case - Operador ──────────────────────────────────────────

def draw_use_case_operador():
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(7, 8.6, "Diagrama de Casos de Uso — Actor: Operador de Almacén",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    # Actor
    ax.plot(1.5, 4.5, 'o', markersize=20, color=GREEN, zorder=5)
    ax.plot([1.5, 1.5], [4.0, 3.2], color=GREEN, lw=2, zorder=5)
    ax.plot([1.0, 2.0], [3.8, 3.8], color=GREEN, lw=2, zorder=5)
    ax.plot([1.5, 1.1], [3.2, 2.5], color=GREEN, lw=2, zorder=5)
    ax.plot([1.5, 1.9], [3.2, 2.5], color=GREEN, lw=2, zorder=5)
    ax.text(1.5, 2.1, "Operador", fontsize=11, ha='center', fontweight='bold', color=GREEN)

    cases = [
        (7, 7.8,  "UC-O1: Escanear QR\nde Reserva", LGREEN, GREEN),
        (7, 6.7,  "UC-O2: Registrar Check-in\n(Almacenar Equipaje)", LBLUE, BLUE),
        (7, 5.6,  "UC-O3: Subir Fotos\nde Evidencia", LORANGE, ORANGE),
        (7, 4.5,  "UC-O4: Etiquetar\nEquipaje", LTEAL, TEAL),
        (7, 3.4,  "UC-O5: Marcar Listo\npara Recojo", LGREEN, GREEN),
        (7, 2.3,  "UC-O6: Solicitar Aprobación\nde Entrega", LPURPLE, PURPLE),
        (7, 1.2,  "UC-O7: Completar Entrega\n(Check-out)", LBLUE, BLUE),
        (11.5, 7.8, "UC-O8: Aprobar/Rechazar\nPago Efectivo", LORANGE, ORANGE),
        (11.5, 6.7, "UC-O9: Verificar Identidad\nde Recojo", LRED, RED),
        (11.5, 5.6, "UC-O10: Registrar Revisión\nde Equipaje", LAMBER, AMBER),
        (11.5, 4.5, "UC-O11: Ver Aprobaciones\nPendientes", LGRAY, GRAY),
        (11.5, 3.4, "UC-O12: Gestionar\nInventario", LTEAL, TEAL),
    ]

    for (cx, cy, label, bg, border) in cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.8, 0.85, facecolor=bg,
                                    edgecolor=border, linewidth=1.8, zorder=3)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, fontsize=7.5, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.15)
        ax.annotate("", xy=(cx - 1.9, cy), xytext=(2.2, min(max(cy, 2.5), 4.5)),
                     arrowprops=dict(arrowstyle='->', color='#A5D6A7', lw=1.0),
                     zorder=2)

    rect = mpatches.FancyBboxPatch((4.5, 0.6), 9.2, 7.8, boxstyle="round,pad=0.3",
                                    facecolor='none', edgecolor=GREEN, linewidth=2,
                                    linestyle='--', zorder=1)
    ax.add_patch(rect)
    ax.text(9.1, 0.8, "Sistema TravelBox Peru — Operaciones QR Handoff", fontsize=9,
            ha='center', color=GREEN, fontstyle='italic')

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "uc_operador.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 3: Use Case - Administrador ─────────────────────────────────────

def draw_use_case_admin():
    fig, ax = plt.subplots(1, 1, figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(7, 8.6, "Diagrama de Casos de Uso — Actor: Administrador",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    # Actor
    ax.plot(1.5, 4.5, 'o', markersize=20, color=RED, zorder=5)
    ax.plot([1.5, 1.5], [4.0, 3.2], color=RED, lw=2, zorder=5)
    ax.plot([1.0, 2.0], [3.8, 3.8], color=RED, lw=2, zorder=5)
    ax.plot([1.5, 1.1], [3.2, 2.5], color=RED, lw=2, zorder=5)
    ax.plot([1.5, 1.9], [3.2, 2.5], color=RED, lw=2, zorder=5)
    ax.text(1.5, 2.1, "Admin", fontsize=11, ha='center', fontweight='bold', color=RED)

    cases = [
        (7, 7.8,  "UC-A1: Gestionar\nUsuarios (CRUD)", LRED, RED),
        (7, 6.7,  "UC-A2: Gestionar\nAlmacenes (CRUD)", LORANGE, ORANGE),
        (7, 5.6,  "UC-A3: Ver Dashboard\n(KPIs / Rankings)", LBLUE, BLUE),
        (7, 4.5,  "UC-A4: Ver Reportes\nde Ingresos", LGREEN, GREEN),
        (7, 3.4,  "UC-A5: Gestionar\nReservas (Bulk)", LPURPLE, PURPLE),
        (7, 2.3,  "UC-A6: Exportar\nDatos (Excel)", LTEAL, TEAL),
        (7, 1.2,  "UC-A7: Gestionar\nCalificaciones", LAMBER, AMBER),
        (11.5, 7.8, "UC-A8: Ver Log\nde Auditoría", LGRAY, GRAY),
        (11.5, 6.7, "UC-A9: Gestionar\nIncidentes", LRED, RED),
        (11.5, 5.6, "UC-A10: Ver Recursos\nAzure", LBLUE, BLUE),
        (11.5, 4.5, "UC-A11: Invalidar\nCache Dashboard", LORANGE, ORANGE),
        (11.5, 3.4, "UC-A12: Gestionar\nReportes i18n", LPURPLE, PURPLE),
    ]

    for (cx, cy, label, bg, border) in cases:
        ellipse = mpatches.Ellipse((cx, cy), 3.8, 0.85, facecolor=bg,
                                    edgecolor=border, linewidth=1.8, zorder=3)
        ax.add_patch(ellipse)
        ax.text(cx, cy, label, fontsize=7.5, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.15)
        ax.annotate("", xy=(cx - 1.9, cy), xytext=(2.2, min(max(cy, 2.5), 4.5)),
                     arrowprops=dict(arrowstyle='->', color='#EF9A9A', lw=1.0),
                     zorder=2)

    rect = mpatches.FancyBboxPatch((4.5, 0.6), 9.2, 7.8, boxstyle="round,pad=0.3",
                                    facecolor='none', edgecolor=RED, linewidth=2,
                                    linestyle='--', zorder=1)
    ax.add_patch(rect)
    ax.text(9.1, 0.8, "Sistema TravelBox Peru — Panel Administrador", fontsize=9,
            ha='center', color=RED, fontstyle='italic')

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "uc_admin.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 4: Reservation State Machine ────────────────────────────────────

def draw_state_machine():
    fig, ax = plt.subplots(1, 1, figsize=(16, 10))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 10)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(8, 9.6, "Diagrama de Estados — Ciclo de Vida de una Reserva",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    states = {
        "DRAFT":            (2.5,  8.0,  LGRAY,   GRAY),
        "PENDING\nPAYMENT": (6.0,  8.0,  LAMBER,  AMBER),
        "CONFIRMED":        (10.0, 8.0,  LGREEN,  GREEN),
        "CHECKIN\nPENDING":  (13.5, 8.0,  LBLUE,   BLUE),
        "STORED":           (13.5, 5.5,  LTEAL,   TEAL),
        "READY FOR\nPICKUP": (10.0, 5.5,  LGREEN,  GREEN),
        "OUT FOR\nDELIVERY": (6.0,  5.5,  LPURPLE, PURPLE),
        "COMPLETED":        (3.0,  3.0,  "#C8E6C9", "#1B5E20"),
        "CANCELLED":        (8.0,  1.5,  LRED,    RED),
        "EXPIRED":          (10.5, 1.5,  LGRAY,   GRAY),
        "INCIDENT":         (13.5, 3.0,  LORANGE, ORANGE),
    }

    # Draw states
    for name, (x, y, bg, border) in states.items():
        box = FancyBboxPatch((x - 1.3, y - 0.45), 2.6, 0.9,
                              boxstyle="round,pad=0.15", facecolor=bg,
                              edgecolor=border, linewidth=2.2, zorder=3)
        ax.add_patch(box)
        ax.text(x, y, name, fontsize=8, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.1)

    # Transitions
    transitions = [
        ("DRAFT",      "PENDING\nPAYMENT", "Crear Reserva"),
        ("PENDING\nPAYMENT", "CONFIRMED",  "Pago OK"),
        ("CONFIRMED",  "CHECKIN\nPENDING",  "Turista llega"),
        ("CHECKIN\nPENDING", "STORED",      "Check-in"),
        ("STORED",     "READY FOR\nPICKUP", "Operador marca"),
        ("READY FOR\nPICKUP", "OUT FOR\nDELIVERY", "Delivery"),
        ("READY FOR\nPICKUP", "COMPLETED",  "Recojo"),
        ("OUT FOR\nDELIVERY", "COMPLETED",  "Entregado"),
        ("STORED",     "INCIDENT",          "Problema"),
        ("CHECKIN\nPENDING", "INCIDENT",    "Problema"),
        ("READY FOR\nPICKUP", "INCIDENT",   "Problema"),
        ("INCIDENT",   "STORED",            "Resuelto"),
        ("INCIDENT",   "COMPLETED",         "Cierre"),
        ("PENDING\nPAYMENT", "CANCELLED",   "Cancela"),
        ("PENDING\nPAYMENT", "EXPIRED",     "Timeout"),
        ("DRAFT",      "CANCELLED",         "Cancela"),
    ]

    drawn_count = {}
    for (src, dst, label) in transitions:
        sx, sy = states[src][0], states[src][1]
        dx, dy = states[dst][0], states[dst][1]

        key = (src, dst) if src < dst else (dst, src)
        drawn_count[key] = drawn_count.get(key, 0) + 1
        offset = 0.15 * (drawn_count[key] - 1)

        mid_x = (sx + dx) / 2
        mid_y = (sy + dy) / 2

        # Determine arrow direction offsets
        angle = math.atan2(dy - sy, dx - sx)
        perp_x = -math.sin(angle) * offset
        perp_y = math.cos(angle) * offset

        ax.annotate("", xy=(dx, dy), xytext=(sx, sy),
                     arrowprops=dict(arrowstyle='->', color='#78909C', lw=1.2,
                                     connectionstyle="arc3,rad=0.15"),
                     zorder=2)
        ax.text(mid_x + perp_x, mid_y + perp_y + 0.18, label,
                fontsize=6.5, ha='center', va='center', color='#546E7A',
                fontstyle='italic', zorder=5,
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                          edgecolor='#B0BEC5', alpha=0.9))

    # Start marker
    ax.plot(0.6, 8.0, 'o', markersize=14, color='black', zorder=5)
    ax.annotate("", xy=(1.2, 8.0), xytext=(0.8, 8.0),
                 arrowprops=dict(arrowstyle='->', color='black', lw=2), zorder=5)

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "state_reserva.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 5: Payment Flow ─────────────────────────────────────────────────

def draw_payment_flow():
    fig, ax = plt.subplots(1, 1, figsize=(15, 12))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 12)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(7.5, 11.5, "Diagrama de Flujo — Proceso de Pago Completo",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    def draw_box(x, y, w, h, text, bg, border, fontsize=8):
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                              boxstyle="round,pad=0.12", facecolor=bg,
                              edgecolor=border, linewidth=2, zorder=3)
        ax.add_patch(box)
        ax.text(x, y, text, fontsize=fontsize, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.15)

    def draw_diamond(x, y, text, bg, border):
        diamond = plt.Polygon([(x, y+0.6), (x+1.2, y), (x, y-0.6), (x-1.2, y)],
                               facecolor=bg, edgecolor=border, linewidth=2, zorder=3)
        ax.add_patch(diamond)
        ax.text(x, y, text, fontsize=7, ha='center', va='center',
                fontweight='bold', color=border, zorder=4)

    def arrow(x1, y1, x2, y2, label="", color='#78909C'):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                     arrowprops=dict(arrowstyle='->', color=color, lw=1.5), zorder=2)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx + 0.15, my, label, fontsize=6.5, color='#546E7A',
                    fontstyle='italic', zorder=5)

    # Flow
    draw_box(4, 10.5, 3.2, 0.7, "1. Turista crea\nReserva", LBLUE, BLUE)
    arrow(4, 10.15, 4, 9.55)

    draw_box(4, 9.2, 3.2, 0.7, "2. POST /payments/intents\n(crear Payment Intent)", LGREEN, GREEN)
    arrow(4, 8.85, 4, 8.25)

    draw_diamond(4, 7.8, "Método\nde pago?", LAMBER, AMBER)
    
    # Card path (left)
    arrow(2.8, 7.8, 1.5, 7.8, "Tarjeta")
    draw_box(1.5, 7.0, 2.8, 0.7, "3a. Backend llama\nIzipay CreatePayment\n→ formToken", LORANGE, ORANGE, 7)
    arrow(1.5, 6.65, 1.5, 6.05)
    draw_box(1.5, 5.7, 2.8, 0.7, "4a. Frontend carga\nSDK Krypton con\nformToken + publicKey", LPURPLE, PURPLE, 7)
    arrow(1.5, 5.35, 1.5, 4.75)
    draw_box(1.5, 4.4, 2.8, 0.7, "5a. Turista llena\ntarjeta en form\nKrypton", LBLUE, BLUE, 7)
    arrow(1.5, 4.05, 1.5, 3.45)
    draw_box(1.5, 3.1, 2.8, 0.7, "6a. Izipay webhook\nPOST /webhooks/izipay\n→ Idempotente", LGREEN, GREEN, 7)
    arrow(2.9, 3.1, 5.5, 2.1)

    # Cash path (right)
    arrow(5.2, 7.8, 7.5, 7.8, "Efectivo")
    draw_box(7.5, 7.0, 2.8, 0.7, "3b. Backend marca\nWAIT_FOR_OPERATOR\n(notifica operador)", LTEAL, TEAL, 7)
    arrow(7.5, 6.65, 7.5, 6.05)
    draw_box(7.5, 5.7, 2.8, 0.7, "4b. Turista paga\nen caja del\nalmacén", LGRAY, GRAY, 7)
    arrow(7.5, 5.35, 7.5, 4.75)
    draw_diamond(7.5, 4.3, "Operador\naprueba?", LAMBER, AMBER)
    
    arrow(6.3, 4.3, 5.5, 4.3, "Sí")
    draw_box(4.5, 4.3, 1.8, 0.6, "Confirma\npago", LGREEN, GREEN, 7)
    arrow(4.5, 4.0, 6.5, 2.1)
    
    arrow(8.7, 4.3, 10.0, 4.3, "No")
    draw_box(11.0, 4.3, 1.8, 0.6, "Rechaza\npago", LRED, RED, 7)
    arrow(11.0, 4.0, 11.0, 2.1)
    draw_box(11.0, 1.8, 2.0, 0.6, "FAILED", LRED, RED)

    # Mock path (center)
    arrow(4, 7.2, 4, 6.05, "Mock")
    draw_box(4, 5.7, 2.8, 0.7, "3c. POST /confirm\napproved=true\n(dev/test)", LGRAY, GRAY, 7)
    arrow(4, 5.35, 6.5, 2.1)

    # Final
    draw_box(6.5, 1.8, 2.5, 0.7, "CONFIRMED\nReserva → CONFIRMED", "#C8E6C9", "#1B5E20")
    arrow(6.5, 1.45, 6.5, 0.75)
    draw_box(6.5, 0.5, 2.8, 0.5, "Email + Notificación al Turista", LBLUE, BLUE, 7)

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "flow_pagos.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 6: Architecture Layers ──────────────────────────────────────────

def draw_architecture():
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(7, 9.6, "Diagrama de Arquitectura — Capas del Sistema",
            fontsize=16, fontweight='bold', ha='center', color=DBLUE)

    layers = [
        (7, 8.5, 12, 1.2, "CAPA DE PRESENTACIÓN (Controllers / REST API)\n"
         "AuthController · PaymentController · ReservationController\n"
         "OpsQrHandoffController · DeliveryController · WarehouseController\n"
         "ProfileController · IncidentController · NotificationController",
         LBLUE, BLUE),
        (7, 6.8, 12, 1.2, "CAPA DE APLICACIÓN (Use Cases / Services)\n"
         "AuthService · PaymentService · ReservationService\n"
         "OpsQrHandoffService · EmailDispatchService · NotificationService\n"
         "InventoryService · DeliveryService · WebhookEventInserter",
         LGREEN, GREEN),
        (7, 5.1, 12, 1.2, "CAPA DE DOMINIO (Entities / Value Objects / Enums)\n"
         "User · Reservation · PaymentAttempt · Warehouse\n"
         "DeliveryOrder · Incident · Rating · Notification\n"
         "ReservationStatus · PaymentStatus · Role",
         LORANGE, ORANGE),
        (7, 3.4, 12, 1.2, "CAPA DE INFRAESTRUCTURA (Adapters / Gateways)\n"
         "IzipayGatewayClient · AzureBlobStorageService\n"
         "JPA Repositories · Flyway Migrations\n"
         "SecurityConfig · JWT Filter · RateLimitFilter",
         LPURPLE, PURPLE),
        (7, 1.7, 12, 1.2, "SERVICIOS EXTERNOS\n"
         "PostgreSQL (Azure) · Azure Blob Storage · Azure Key Vault\n"
         "Izipay/Krypton API · Microsoft Graph (Email)\n"
         "Google OAuth 2.0 · Microsoft Entra ID",
         LGRAY, GRAY),
    ]

    for (x, y, w, h, text, bg, border) in layers:
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                              boxstyle="round,pad=0.15", facecolor=bg,
                              edgecolor=border, linewidth=2.5, zorder=3)
        ax.add_patch(box)
        ax.text(x, y, text, fontsize=7.8, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.3)

    # Arrows between layers
    for y_top, y_bot in [(7.9, 7.4), (6.2, 5.7), (4.5, 4.0), (2.8, 2.3)]:
        ax.annotate("", xy=(7, y_bot), xytext=(7, y_top),
                     arrowprops=dict(arrowstyle='->', color='#90A4AE', lw=2.5), zorder=2)

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "architecture.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── DIAGRAM 7: QR Handoff Workflow ───────────────────────────────────────────

def draw_qr_handoff_flow():
    fig, ax = plt.subplots(1, 1, figsize=(14, 12))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 12)
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(7, 11.5, "Diagrama de Flujo — Operación QR Handoff (Check-in a Check-out)",
            fontsize=15, fontweight='bold', ha='center', color=DBLUE)

    def draw_box(x, y, w, h, text, bg, border, fs=8):
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                              boxstyle="round,pad=0.12", facecolor=bg,
                              edgecolor=border, linewidth=2, zorder=3)
        ax.add_patch(box)
        ax.text(x, y, text, fontsize=fs, ha='center', va='center',
                fontweight='bold', color=border, zorder=4, linespacing=1.15)

    def arrow(x1, y1, x2, y2, label=""):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                     arrowprops=dict(arrowstyle='->', color='#78909C', lw=1.5), zorder=2)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx + 0.2, my, label, fontsize=6.5, color='#546E7A', fontstyle='italic')

    steps = [
        (7, 10.7, 4.5, 0.65, "1. Turista llega con QR\n(reserva CONFIRMED)", LBLUE, BLUE),
        (7, 9.7,  4.5, 0.65, "2. Operador escanea QR\nPOST /ops/qr-handoff/scan", LGREEN, GREEN),
        (7, 8.7,  4.5, 0.65, "3. Operador registra almacenamiento\nPOST .../store-with-photos", LORANGE, ORANGE),
        (7, 7.7,  4.5, 0.65, "4. Operador etiqueta equipaje\nPOST .../tag → STORED", LTEAL, TEAL),
        (7, 6.7,  4.5, 0.65, "5. Operador marca listo\nPOST .../ready-for-pickup", LGREEN, GREEN),
        (3.5, 5.5, 3.5, 0.65, "6a. Recojo en almacén:\nVerificar identidad\nPATCH .../delivery/identity", LPURPLE, PURPLE),
        (10.5, 5.5, 3.5, 0.65, "6b. Delivery:\nSolicitar aprobación\nPOST .../request-approval", LAMBER, AMBER),
        (3.5, 4.3,  3.5, 0.65, "7a. Solicitar aprobación\nPOST .../request-approval", LBLUE, BLUE),
        (10.5, 4.3, 3.5, 0.65, "7b. Supervisor aprueba\nPOST /approvals/{id}/approve", LGREEN, GREEN),
        (3.5, 3.1,  3.5, 0.65, "8a. Supervisor aprueba\nPOST /approvals/{id}/approve", LGREEN, GREEN),
        (10.5, 3.1, 3.5, 0.65, "8b. Completar entrega\nPOST .../delivery/complete", LORANGE, ORANGE),
        (7, 1.8,  4.5, 0.65, "9. COMPLETED\nTurista califica almacén", "#C8E6C9", "#1B5E20"),
        (7, 0.8,  4.5, 0.5, "Notificación + Email al Turista", LBLUE, BLUE),
    ]

    for (x, y, w, h, text, bg, border) in steps:
        draw_box(x, y, w, h, text, bg, border, 7)

    # Vertical arrows
    for i in range(4):
        arrow(7, steps[i][1] - 0.33, 7, steps[i+1][1] + 0.33)
    
    # Split to pickup / delivery
    arrow(5.5, 6.37, 3.5, 5.83, "Recojo")
    arrow(8.5, 6.37, 10.5, 5.83, "Delivery")
    
    # Pickup path
    arrow(3.5, 5.17, 3.5, 4.63)
    arrow(3.5, 3.97, 3.5, 3.43)
    arrow(3.5, 2.77, 5.5, 2.13)
    
    # Delivery path
    arrow(10.5, 5.17, 10.5, 4.63)
    arrow(10.5, 3.97, 10.5, 3.43)
    arrow(10.5, 2.77, 8.5, 2.13)
    
    # To completed
    arrow(7, 1.47, 7, 1.05)

    plt.tight_layout()
    path = os.path.join(IMG_DIR, "flow_qr_handoff.png")
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

# ── BUILD THE WORD DOCUMENT ─────────────────────────────────────────────────

def build_word():
    print("Generating diagrams...")
    img_turista    = draw_use_case_turista()
    img_operador   = draw_use_case_operador()
    img_admin      = draw_use_case_admin()
    img_states     = draw_state_machine()
    img_payments   = draw_payment_flow()
    img_arch       = draw_architecture()
    img_qr         = draw_qr_handoff_flow()
    print("All diagrams generated.")

    doc = Document()

    # ── Page setup
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── COVER PAGE ───────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TRAVELBOX PERU")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plataforma de Almacenamiento de Equipaje para Turistas")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

    doc.add_paragraph()

    info_lines = [
        "Diagramas de Casos de Uso",
        "Diagramas de Estado y Flujo",
        "Diagrama de Arquitectura",
        "Datos de Prueba Reales",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {line}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("API Base URL: https://api.inkavoy.pe")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión 1.0 — Abril 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)

    doc.add_page_break()

    # ── TABLE OF CONTENTS (Manual) ───────────────────────────────────────
    add_styled_heading(doc, "Tabla de Contenido", 1)
    toc = [
        "1. Diagrama de Casos de Uso — Turista",
        "2. Diagrama de Casos de Uso — Operador de Almacén",
        "3. Diagrama de Casos de Uso — Administrador",
        "4. Diagrama de Estados — Ciclo de Vida de Reserva",
        "5. Diagrama de Flujo — Proceso de Pago Completo",
        "6. Diagrama de Flujo — Operación QR Handoff",
        "7. Diagrama de Arquitectura — Capas del Sistema",
        "8. Datos de Prueba — Usuarios",
        "9. Datos de Prueba — Almacenes",
        "10. Datos de Prueba — Flujos Completos (E2E con cURL)",
        "11. Datos de Prueba — Tarjetas Izipay Sandbox",
    ]
    for item in toc:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ── 1. USE CASE TURISTA ──────────────────────────────────────────────
    add_styled_heading(doc, "1. Diagrama de Casos de Uso — Turista", 1)
    p = doc.add_paragraph()
    p.add_run("Actor principal: ").bold = True
    p.add_run("Turista / Viajero (usuario registrado con rol USER)")
    doc.add_paragraph()
    doc.add_picture(img_turista, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_heading(doc, "Detalle de Casos de Uso — Turista", 2)

    uc_turista_detail = [
        ("UC-01", "Registrarse / Login", "POST /api/v1/auth/register\nPOST /api/v1/auth/login",
         "El turista crea una cuenta con email+contraseña o inicia sesión. Se envía email de verificación."),
        ("UC-02", "Login Social", "GET /api/v1/auth/oauth/{provider}/start\nGET /api/v1/auth/oauth/{provider}/callback",
         "Login via Google o Microsoft (Entra ID). Redirige al provider y callback registra/autentica."),
        ("UC-03", "Buscar Almacenes Cercanos", "GET /api/v1/warehouses/search\nGET /api/v1/warehouses/nearby\nGET /api/v1/geo/warehouses/nearby",
         "Busca almacenes por ubicación GPS, ciudad, zona. Incluye precio, capacidad y ratings."),
        ("UC-04", "Crear Reserva", "POST /api/v1/reservations",
         "Crea reserva indicando almacén, fechas check-in/out y cantidad de maletas. Estado inicial: DRAFT -> PENDING_PAYMENT."),
        ("UC-05", "Pagar con Tarjeta (Izipay)", "POST /api/v1/payments/intents\n→ formToken + publicKey",
         "Crea Payment Intent, backend llama a Izipay CreatePayment, devuelve formToken para SDK Krypton."),
        ("UC-06", "Pagar en Efectivo", "POST /api/v1/payments/intents\n(paymentMethod=cash)",
         "Crea intent de pago en efectivo. Operador del almacén recibe notificación para aprobar."),
        ("UC-07", "Ver Tracking de Reserva", "GET /api/v1/reservations/{id}/tracking",
         "Consulta estado actual de la reserva con timeline completo de transiciones."),
        ("UC-08", "Recibir Notificaciones", "GET /api/v1/notifications/my\nGET /api/v1/notifications/events (SSE)",
         "Recibe notificaciones push via SSE (Server-Sent Events) y consulta historial."),
        ("UC-09", "Calificar Almacén", "POST /api/v1/ratings",
         "Después de completar reserva, califica el servicio con estrellas y comentario."),
        ("UC-10", "Gestionar Perfil", "GET /api/v1/profile/me\nPATCH /api/v1/profile/me",
         "Ver y actualizar datos personales, subir foto de perfil."),
        ("UC-11", "Cambiar Email", "POST /api/v1/auth/email-change/initiate\nPOST /api/v1/auth/email-change/verify",
         "Solicita cambio de email con verificación por código enviado al nuevo email."),
        ("UC-12", "Recuperar Contraseña", "POST /api/v1/auth/password-reset/request\nPOST /api/v1/auth/password-reset/confirm",
         "Flujo de reset password por email con token de un solo uso."),
        ("UC-13", "Reportar Incidente", "POST /api/v1/incidents\nPOST /api/v1/incidents/{id}/messages",
         "Reporta problema con equipaje. Chat de mensajes entre turista y soporte."),
        ("UC-14", "Ver Historial de Pagos", "GET /api/v1/payments/history",
         "Consulta todos los pagos realizados con estado, método y fecha."),
        ("UC-15", "Cancelar Reserva", "PATCH /api/v1/reservations/{id}/cancel",
         "Cancela reserva si el estado lo permite (DRAFT, PENDING_PAYMENT, CONFIRMED)."),
    ]

    for (code, name, endpoints, desc) in uc_turista_detail:
        table = doc.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cells = table.rows[0].cells
        cells[0].text = "Código"
        cells[1].text = code
        set_cell_bg(cells[0], LBLUE.replace("#", ""))

        cells = table.rows[1].cells
        cells[0].text = "Caso de Uso"
        cells[1].text = name
        set_cell_bg(cells[0], LBLUE.replace("#", ""))

        cells = table.rows[2].cells
        cells[0].text = "Endpoints"
        cells[1].text = endpoints
        set_cell_bg(cells[0], LBLUE.replace("#", ""))

        cells = table.rows[3].cells
        cells[0].text = "Descripción"
        cells[1].text = desc
        set_cell_bg(cells[0], LBLUE.replace("#", ""))

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(5.3)

        doc.add_paragraph()  # spacing

    doc.add_page_break()

    # ── 2. USE CASE OPERADOR ─────────────────────────────────────────────
    add_styled_heading(doc, "2. Diagrama de Casos de Uso — Operador de Almacén", 1)
    p = doc.add_paragraph()
    p.add_run("Actor: ").bold = True
    p.add_run("Operador de Almacén (rol OPERATOR) — Personal que recibe, almacena y entrega equipaje.")
    doc.add_paragraph()
    doc.add_picture(img_operador, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_heading(doc, "Detalle de Casos de Uso — Operador", 2)

    uc_op = [
        ("UC-O1", "Escanear QR de Reserva", "POST /api/v1/ops/qr-handoff/scan",
         "Escanea código QR de la reserva del turista. Valida que la reserva esté en estado correcto."),
        ("UC-O2", "Registrar Check-in", "POST .../store-with-photos",
         "Registra el almacenamiento del equipaje con fotos de evidencia. Estado → STORED."),
        ("UC-O3", "Subir Fotos de Evidencia", "POST /api/v1/inventory/evidences/upload",
         "Sube fotos del equipaje a Azure Blob Storage como evidencia."),
        ("UC-O4", "Etiquetar Equipaje", "POST .../tag",
         "Asigna etiqueta/tag físico al equipaje almacenado."),
        ("UC-O5", "Marcar Listo para Recojo", "POST .../ready-for-pickup",
         "Marca la reserva como lista. Estado → READY_FOR_PICKUP. Notifica al turista."),
        ("UC-O6", "Solicitar Aprobación de Entrega", "POST .../request-approval",
         "Solicita a un supervisor la aprobación para entregar el equipaje."),
        ("UC-O7", "Completar Entrega (Check-out)", "POST .../delivery/complete",
         "Confirma la entrega del equipaje al turista. Estado → COMPLETED."),
        ("UC-O8", "Aprobar/Rechazar Pago Efectivo", "POST /payments/cash/{id}/approve\nPOST /payments/cash/{id}/reject",
         "Confirma o rechaza que el turista pagó en efectivo en la caja del almacén."),
        ("UC-O9", "Verificar Identidad de Recojo", "PATCH .../delivery/identity",
         "Verifica la identidad del turista o tercero autorizado al momento de la entrega."),
        ("UC-O10", "Registrar Revisión de Equipaje", "PATCH .../delivery/luggage",
         "Registra el estado del equipaje antes de la entrega (fotos, observaciones)."),
        ("UC-O11", "Ver Aprobaciones Pendientes", "GET /api/v1/ops/qr-handoff/approvals",
         "Lista todas las solicitudes de aprobación pendientes para su almacén."),
        ("UC-O12", "Gestionar Inventario", "POST /api/v1/inventory/checkin\nPOST /api/v1/inventory/checkout",
         "Registra entradas y salidas de equipaje del inventario del almacén."),
    ]

    for (code, name, endpoints, desc) in uc_op:
        table = doc.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (lbl, val) in enumerate([("Código", code), ("Caso de Uso", name),
                                          ("Endpoints", endpoints), ("Descripción", desc)]):
            table.rows[i].cells[0].text = lbl
            table.rows[i].cells[1].text = val
            set_cell_bg(table.rows[i].cells[0], LGREEN.replace("#", ""))
            for paragraph in table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            for paragraph in table.rows[i].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            table.rows[i].cells[0].width = Inches(1.2)
            table.rows[i].cells[1].width = Inches(5.3)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 3. USE CASE ADMIN ────────────────────────────────────────────────
    add_styled_heading(doc, "3. Diagrama de Casos de Uso — Administrador", 1)
    p = doc.add_paragraph()
    p.add_run("Actor: ").bold = True
    p.add_run("Administrador del Sistema (rol ADMIN) — Gestión total de la plataforma.")
    doc.add_paragraph()
    doc.add_picture(img_admin, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_heading(doc, "Detalle de Casos de Uso — Administrador", 2)

    uc_admin = [
        ("UC-A1", "Gestionar Usuarios (CRUD)", "GET/POST/PUT/DELETE /api/v1/admin/users",
         "CRUD completo de usuarios. Cambiar roles, activar/desactivar, reset password, bulk operations."),
        ("UC-A2", "Gestionar Almacenes (CRUD)", "GET/POST/PUT/DELETE /api/v1/admin/warehouses",
         "CRUD de almacenes. Subir fotos, editar precios, capacidad, ubicación GPS."),
        ("UC-A3", "Ver Dashboard (KPIs)", "GET /api/v1/admin/dashboard",
         "Resumen ejecutivo: reservas activas, ingresos, top operadores, tendencias."),
        ("UC-A4", "Ver Reportes de Ingresos", "GET /api/v1/admin/reports/revenue",
         "Reporte de ingresos por período, almacén y método de pago."),
        ("UC-A5", "Gestionar Reservas (Bulk)", "PATCH /api/v1/admin/reservations/bulk/status",
         "Operaciones masivas sobre reservas: cambiar estado, cancelar lotes."),
        ("UC-A6", "Exportar Datos (Excel)", "GET /api/v1/admin/users/export\nGET /api/v1/admin/reservations/export",
         "Exporta usuarios y reservas en formato Excel (.xlsx)."),
        ("UC-A7", "Gestionar Calificaciones", "GET/PATCH/DELETE /api/v1/admin/reports/ratings",
         "Moderar reviews: ocultar, editar, eliminar calificaciones inapropiadas."),
        ("UC-A8", "Ver Log de Auditoría", "GET /api/v1/admin/system/audit-log",
         "Consulta bitácora de acciones administrativas realizadas."),
        ("UC-A9", "Gestionar Incidentes", "GET /api/v1/incidents/page\nPATCH /api/v1/incidents/{id}/resolve",
         "Ver y resolver incidentes reportados por turistas."),
        ("UC-A10", "Ver Recursos Azure", "GET /api/v1/admin/system/azure-resources",
         "Consulta estado de recursos Azure: DB, storage, app service."),
        ("UC-A11", "Invalidar Cache Dashboard", "POST /api/v1/admin/dashboard/invalidate-cache",
         "Fuerza recálculo del dashboard (útil después de correcciones manuales)."),
        ("UC-A12", "Gestionar Reportes i18n", "GET/POST/DELETE /api/v1/admin/i18n-report",
         "Ver, crear y limpiar reportes de errores de traducción/internacionalización."),
    ]

    for (code, name, endpoints, desc) in uc_admin:
        table = doc.add_table(rows=4, cols=2, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (lbl, val) in enumerate([("Código", code), ("Caso de Uso", name),
                                          ("Endpoints", endpoints), ("Descripción", desc)]):
            table.rows[i].cells[0].text = lbl
            table.rows[i].cells[1].text = val
            set_cell_bg(table.rows[i].cells[0], LRED.replace("#", ""))
            for paragraph in table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            for paragraph in table.rows[i].cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            table.rows[i].cells[0].width = Inches(1.2)
            table.rows[i].cells[1].width = Inches(5.3)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 4. STATE DIAGRAM ─────────────────────────────────────────────────
    add_styled_heading(doc, "4. Diagrama de Estados — Ciclo de Vida de una Reserva", 1)
    p = doc.add_paragraph()
    p.add_run("Máquina de estados definida en: ").bold = True
    p.add_run("ReservationStatus.java — Todas las transiciones están validadas por el backend.")
    doc.add_paragraph()
    doc.add_picture(img_states, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_heading(doc, "Tabla de Transiciones de Estado", 2)

    transitions_table_data = [
        ("DRAFT", "PENDING_PAYMENT", "Turista envía reserva"),
        ("DRAFT", "CANCELLED", "Turista cancela borrador"),
        ("PENDING_PAYMENT", "CONFIRMED", "Pago aprobado (tarjeta/efectivo/mock)"),
        ("PENDING_PAYMENT", "CANCELLED", "Turista cancela"),
        ("PENDING_PAYMENT", "EXPIRED", "Timeout sin pago"),
        ("CONFIRMED", "CHECKIN_PENDING", "Turista llega al almacén"),
        ("CONFIRMED", "CANCELLED", "Cancela antes de check-in"),
        ("CHECKIN_PENDING", "STORED", "Operador almacena equipaje"),
        ("CHECKIN_PENDING", "INCIDENT", "Problema detectado"),
        ("CHECKIN_PENDING", "CANCELLED", "Cancela en check-in"),
        ("STORED", "READY_FOR_PICKUP", "Operador marca listo"),
        ("STORED", "OUT_FOR_DELIVERY", "Inicia delivery"),
        ("STORED", "INCIDENT", "Problema con equipaje"),
        ("READY_FOR_PICKUP", "COMPLETED", "Turista recoge equipaje"),
        ("READY_FOR_PICKUP", "OUT_FOR_DELIVERY", "Cambio a delivery"),
        ("READY_FOR_PICKUP", "INCIDENT", "Problema en entrega"),
        ("OUT_FOR_DELIVERY", "COMPLETED", "Entrega exitosa"),
        ("OUT_FOR_DELIVERY", "INCIDENT", "Problema en delivery"),
        ("INCIDENT", "STORED", "Resuelto, vuelve a almacén"),
        ("INCIDENT", "READY_FOR_PICKUP", "Resuelto, listo para recojo"),
        ("INCIDENT", "COMPLETED", "Cierre administrativo"),
        ("INCIDENT", "CANCELLED", "Cancelación por incidente"),
    ]

    t = doc.add_table(rows=len(transitions_table_data) + 1, cols=3, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Estado Origen", "Estado Destino", "Disparador"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        set_cell_bg(t.rows[0].cells[i], DBLUE.replace("#", ""))
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    for idx, (src, dst, trigger) in enumerate(transitions_table_data):
        row = t.rows[idx + 1]
        row.cells[0].text = src
        row.cells[1].text = dst
        row.cells[2].text = trigger
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_page_break()

    # ── 5. PAYMENT FLOW ──────────────────────────────────────────────────
    add_styled_heading(doc, "5. Diagrama de Flujo — Proceso de Pago Completo", 1)
    p = doc.add_paragraph()
    p.add_run("Tres caminos de pago: ").bold = True
    p.add_run("Tarjeta (Izipay Krypton SDK), Efectivo (aprobación por operador), y Mock (pruebas).")
    doc.add_paragraph()
    doc.add_picture(img_payments, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── 6. QR HANDOFF FLOW ───────────────────────────────────────────────
    add_styled_heading(doc, "6. Diagrama de Flujo — Operación QR Handoff", 1)
    p = doc.add_paragraph()
    p.add_run("Flujo completo desde que el turista llega con su QR hasta la entrega final del equipaje.")
    doc.add_paragraph()
    doc.add_picture(img_qr, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── 7. ARCHITECTURE ──────────────────────────────────────────────────
    add_styled_heading(doc, "7. Diagrama de Arquitectura — Capas del Sistema", 1)
    p = doc.add_paragraph()
    p.add_run("Arquitectura Hexagonal (Clean Architecture) con Spring Boot 3.5 / Java 21.")
    doc.add_paragraph()
    doc.add_picture(img_arch, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── 8. TEST DATA - USERS ─────────────────────────────────────────────
    add_styled_heading(doc, "8. Datos de Prueba — Usuarios", 1)
    p = doc.add_paragraph()
    p.add_run("Usuarios pre-cargados en el entorno de producción/sandbox. ").bold = True
    p.add_run("Usar estos datos contra: https://api.inkavoy.pe")

    doc.add_paragraph()

    users_data = [
        ("admin@travelbox.pe", "Admin123!", "ADMIN", "Administrador del sistema"),
        ("operator@travelbox.pe", "Operator123!", "OPERATOR", "Operador de almacén"),
        ("turista1@gmail.com", "Tourist123!", "USER", "Turista de prueba 1"),
        ("turista2@gmail.com", "Tourist123!", "USER", "Turista de prueba 2"),
        ("supervisor@travelbox.pe", "Super123!", "CITY_SUPERVISOR", "Supervisor de zona"),
    ]

    t = doc.add_table(rows=len(users_data) + 1, cols=4, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Email", "Contraseña", "Rol", "Descripción"]):
        t.rows[0].cells[i].text = h
        set_cell_bg(t.rows[0].cells[i], DBLUE.replace("#", ""))
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    for idx, (email, pwd, role, desc) in enumerate(users_data):
        row = t.rows[idx + 1]
        row.cells[0].text = email
        row.cells[1].text = pwd
        row.cells[2].text = role
        row.cells[3].text = desc
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_separator(doc)

    # ── 9. TEST DATA - WAREHOUSES ────────────────────────────────────────
    add_styled_heading(doc, "9. Datos de Prueba — Almacenes", 1)

    warehouses_data = [
        ("Almacén Centro Lima", "Lima", "Cercado de Lima", "-12.0464", "-77.0428", "S/. 15.00 /día", "50 maletas"),
        ("Almacén Miraflores", "Lima", "Miraflores", "-12.1191", "-77.0300", "S/. 20.00 /día", "30 maletas"),
        ("Almacén Cusco Plaza", "Cusco", "Centro Histórico", "-13.5170", "-71.9785", "S/. 12.00 /día", "40 maletas"),
        ("Almacén Arequipa", "Arequipa", "Centro", "-16.3989", "-71.5370", "S/. 10.00 /día", "25 maletas"),
        ("Almacén Aeropuerto Jorge Chávez", "Lima", "Callao", "-12.0219", "-77.1143", "S/. 25.00 /día", "100 maletas"),
    ]

    t = doc.add_table(rows=len(warehouses_data) + 1, cols=7, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Nombre", "Ciudad", "Zona", "Latitud", "Longitud", "Precio", "Capacidad"]):
        t.rows[0].cells[i].text = h
        set_cell_bg(t.rows[0].cells[i], GREEN.replace("#", ""))
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    for idx, row_data in enumerate(warehouses_data):
        row = t.rows[idx + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_page_break()

    # ── 10. TEST DATA - FULL E2E FLOWS ───────────────────────────────────
    add_styled_heading(doc, "10. Datos de Prueba — Flujos Completos (E2E)", 1)
    p = doc.add_paragraph()
    p.add_run("Todos los comandos usan la URL base: ").bold = True
    p.add_run("https://api.inkavoy.pe")
    doc.add_paragraph()

    add_styled_heading(doc, "Flujo A: Registro + Reserva + Pago con Tarjeta (Izipay)", 2)

    flows_card = [
        ("Paso 1: Registrar cuenta",
         'POST /api/v1/auth/register\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "email": "turista.test@gmail.com",\n'
         '  "password": "MiPassword123!",\n'
         '  "firstName": "Carlos",\n'
         '  "lastName": "Méndez",\n'
         '  "phone": "+51999888777"\n'
         '}'),
        ("Paso 2: Login",
         'POST /api/v1/auth/login\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "email": "turista.test@gmail.com",\n'
         '  "password": "MiPassword123!"\n'
         '}\n\n'
         '→ Respuesta: { "accessToken": "eyJhbG...", "refreshToken": "..." }'),
        ("Paso 3: Buscar almacenes cercanos",
         'GET /api/v1/warehouses/nearby?lat=-12.0464&lon=-77.0428&radiusKm=5\n'
         'Authorization: Bearer {accessToken}'),
        ("Paso 4: Crear reserva",
         'POST /api/v1/reservations\n'
         'Authorization: Bearer {accessToken}\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "warehouseId": 1,\n'
         '  "baggageCount": 2,\n'
         '  "checkInDate": "2026-04-05T10:00:00",\n'
         '  "checkOutDate": "2026-04-07T18:00:00",\n'
         '  "notes": "2 maletas grandes"\n'
         '}\n\n'
         '→ Respuesta: { "id": 42, "status": "PENDING_PAYMENT", ... }'),
        ("Paso 5: Crear Payment Intent (tarjeta)",
         'POST /api/v1/payments/intents\n'
         'Authorization: Bearer {accessToken}\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "reservationId": 42,\n'
         '  "paymentMethod": "card"\n'
         '}\n\n'
         '→ Respuesta incluye nextAction:\n'
         '  "authorization": "<formToken>",\n'
         '  "keyRSA": "23662223:testpublickey...",\n'
         '  "scriptUrl": "https://static.micuentaweb.pe/...kr-payment-form.min.js"'),
        ("Paso 6: En el frontend cargar SDK Krypton",
         'El frontend inyecta el script de nextAction.scriptUrl,\n'
         'configura kr-public-key con nextAction.keyRSA,\n'
         'y pasa kr-form-token con nextAction.authorization.\n\n'
         'Tarjeta de prueba: 4970 1000 0000 0055\n'
         'CVV: 123 | Vencimiento: cualquier fecha futura'),
        ("Paso 7: Webhook automático de Izipay",
         'Izipay llama automáticamente:\n'
         'POST /api/v1/payments/webhooks/izipay\n\n'
         '→ Backend actualiza PaymentAttempt.status = CONFIRMED\n'
         '→ Reservation.status = CONFIRMED\n'
         '→ Email + notificación al turista'),
    ]

    for (title, body) in flows_card:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(8.5)
        run.font.name = "Consolas"
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    add_separator(doc)

    add_styled_heading(doc, "Flujo B: Pago en Efectivo", 2)

    flows_cash = [
        ("Paso 1-4: Igual que Flujo A (registro, login, buscar, crear reserva)",
         "Mismos pasos anteriores hasta tener reservationId."),
        ("Paso 5: Crear Payment Intent (efectivo)",
         'POST /api/v1/payments/intents\n'
         'Authorization: Bearer {accessToken}\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "reservationId": 42,\n'
         '  "paymentMethod": "cash"\n'
         '}\n\n'
         '→ nextAction.type = "WAIT_FOR_OPERATOR"\n'
         '→ Operador del almacén recibe notificación'),
        ("Paso 6: Operador aprueba pago (con token de operador)",
         'POST /api/v1/payments/cash/{paymentIntentId}/approve\n'
         'Authorization: Bearer {operatorToken}\n'
         'Content-Type: application/json\n\n'
         '{\n'
         '  "reason": "Turista pagó S/. 30 en caja"\n'
         '}\n\n'
         '→ PaymentAttempt.status = CONFIRMED\n'
         '→ Reservation.status = CONFIRMED'),
    ]

    for (title, body) in flows_cash:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(8.5)
        run.font.name = "Consolas"
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    add_separator(doc)

    add_styled_heading(doc, "Flujo C: Pago Mock (para desarrollo/pruebas sin Izipay)", 2)

    p = doc.add_paragraph()
    run = p.add_run("Paso 5 alternativo: Confirmar pago mock")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

    p = doc.add_paragraph()
    run = p.add_run(
        'POST /api/v1/payments/confirm\n'
        'Authorization: Bearer {accessToken}\n'
        'Content-Type: application/json\n\n'
        '{\n'
        '  "paymentIntentId": 42,\n'
        '  "paymentMethod": "card",\n'
        '  "approved": true\n'
        '}\n\n'
        '→ Funciona porque app.payments.allow-mock-confirmation=true\n'
        '→ Pago confirmado inmediatamente sin pasarela real'
    )
    run.font.size = Pt(8.5)
    run.font.name = "Consolas"
    p.paragraph_format.left_indent = Cm(0.5)

    doc.add_page_break()

    add_styled_heading(doc, "Flujo D: Operación QR Handoff completa", 2)

    flows_qr = [
        ("Paso 1: Login como operador",
         'POST /api/v1/auth/login\n'
         '{ "email": "operator@travelbox.pe", "password": "Operator123!" }'),
        ("Paso 2: Escanear QR",
         'POST /api/v1/ops/qr-handoff/scan\n'
         'Authorization: Bearer {operatorToken}\n'
         '{ "qrCode": "TBOX-RES-42" }'),
        ("Paso 3: Almacenar equipaje con fotos",
         'POST /api/v1/ops/qr-handoff/reservations/42/store-with-photos\n'
         'Authorization: Bearer {operatorToken}\n'
         'Content-Type: multipart/form-data\n'
         'photos: [file1.jpg, file2.jpg]\n'
         'notes: "2 maletas grandes, buen estado"'),
        ("Paso 4: Etiquetar",
         'POST /api/v1/ops/qr-handoff/reservations/42/tag\n'
         'Authorization: Bearer {operatorToken}\n'
         '{ "tagCode": "TAG-A42-001" }'),
        ("Paso 5: Marcar listo para recojo",
         'POST /api/v1/ops/qr-handoff/reservations/42/ready-for-pickup\n'
         'Authorization: Bearer {operatorToken}'),
        ("Paso 6: Confirmar recojo del turista",
         'POST /api/v1/ops/qr-handoff/reservations/42/pickup/confirm\n'
         'Authorization: Bearer {operatorToken}'),
        ("Paso 7: Solicitar aprobación de entrega",
         'POST /api/v1/ops/qr-handoff/reservations/42/delivery/request-approval\n'
         'Authorization: Bearer {operatorToken}'),
        ("Paso 8: Supervisor aprueba",
         'POST /api/v1/ops/qr-handoff/approvals/{approvalId}/approve\n'
         'Authorization: Bearer {supervisorToken}'),
        ("Paso 9: Completar entrega",
         'POST /api/v1/ops/qr-handoff/reservations/42/delivery/complete\n'
         'Authorization: Bearer {operatorToken}\n\n'
         '→ Reserva pasa a COMPLETED\n'
         '→ Turista recibe notificación + email'),
    ]

    for (title, body) in flows_qr:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x69, 0x5C)

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(8.5)
        run.font.name = "Consolas"
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 11. IZIPAY TEST CARDS ────────────────────────────────────────────
    add_styled_heading(doc, "11. Datos de Prueba — Tarjetas Izipay Sandbox", 1)
    p = doc.add_paragraph()
    p.add_run("Estas tarjetas solo funcionan con credenciales de prueba (merchant: 23662223).")
    doc.add_paragraph()

    cards = [
        ("VISA Aprobada", "4970 1000 0000 0055", "123", "Cualquier fecha futura", "PAID"),
        ("VISA Rechazada", "4970 1000 0000 0014", "123", "Cualquier fecha futura", "REFUSED"),
        ("MASTERCARD Aprobada", "5100 0100 0000 0016", "123", "Cualquier fecha futura", "PAID"),
        ("MASTERCARD Rechazada", "5100 0100 0000 0024", "123", "Cualquier fecha futura", "REFUSED"),
        ("VISA 3DS Challenge", "4970 1000 0000 0048", "123", "Cualquier fecha futura", "PAID (con 3DS)"),
    ]

    t = doc.add_table(rows=len(cards) + 1, cols=5, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Tipo", "Número Tarjeta", "CVV", "Vencimiento", "Resultado"]):
        t.rows[0].cells[i].text = h
        set_cell_bg(t.rows[0].cells[i], ORANGE.replace("#", ""))
        for paragraph in t.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    for idx, row_data in enumerate(cards):
        row = t.rows[idx + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Documento generado automáticamente para TravelBox Peru —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)
    run.font.italic = True

    # ── SAVE ─────────────────────────────────────────────────────────────
    output_path = os.path.join(OUT_DIR, "TravelBox_Peru_Documentacion_Completa.docx")
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print(f"Diagrams saved to: {IMG_DIR}")
    return output_path


if __name__ == "__main__":
    build_word()
